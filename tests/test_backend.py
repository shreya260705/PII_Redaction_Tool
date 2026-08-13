import os
import io
import time
import docx
import pytest
from fastapi.testclient import TestClient

from backend.main import app, file_mappings, TEMP_BASE_DIR

client = TestClient(app)

# Helper to create a valid minimal DOCX in memory
def create_minimal_docx() -> bytes:
    doc = docx.Document()
    doc.add_paragraph("Hello, this is a test document with Rashi Patil and rashi@gmail.com.")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def test_health_endpoint():
    response = client.get("/api/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}

def test_valid_docx_upload():
    docx_bytes = create_minimal_docx()
    response = client.post(
        "/api/redact",
        files={"file": ("test_doc.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 200
    data = response.json()
    assert data["success"] is True
    assert "file_id" in data
    assert "total_replacements" in data
    assert "replacements_by_type" in data
    assert "unique_mappings_count" in data
    # Ensure no original PII is leaked in the metadata response
    assert "Rashi" not in str(data)
    assert "rashi@gmail.com" not in str(data)

def test_invalid_extension():
    response = client.post(
        "/api/redact",
        files={"file": ("test_doc.txt", b"dummy content", "text/plain")}
    )
    assert response.status_code == 400
    assert "Only .docx files are allowed" in response.json()["detail"]

def test_large_file_limit():
    # 20 MB + 100 bytes
    large_bytes = b"0" * (20 * 1024 * 1024 + 100)
    response = client.post(
        "/api/redact",
        files={"file": ("large_doc.docx", large_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert response.status_code == 400
    assert "File too large" in response.json()["detail"]

def test_download_workflow():
    docx_bytes = create_minimal_docx()
    # 1. Upload and redact
    redact_resp = client.post(
        "/api/redact",
        files={"file": ("my_doc.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    assert redact_resp.status_code == 200
    file_id = redact_resp.json()["file_id"]

    # 2. Download file
    download_resp = client.get(f"/api/download/{file_id}")
    assert download_resp.status_code == 200
    assert download_resp.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    # 3. Verify downloaded DOCX opens successfully
    downloaded_bytes = download_resp.content
    bio = io.BytesIO(downloaded_bytes)
    doc = docx.Document(bio)
    assert len(doc.paragraphs) == 1
    # Check that original PII is absent
    assert "Rashi Patil" not in doc.paragraphs[0].text
    assert "rashi@gmail.com" not in doc.paragraphs[0].text

def test_invalid_file_id_returns_404():
    # Invalid UUID format
    response = client.get("/api/download/nonexistent-id")
    assert response.status_code == 404
    # Valid UUID format but nonexistent
    response = client.get("/api/download/00000000-0000-0000-0000-000000000000")
    assert response.status_code == 404

def test_temporary_cleanup():
    # Add dummy mapping with creation time in the past (>30 mins ago)
    dummy_id = "11111111-1111-1111-1111-111111111111"
    temp_file = os.path.join(TEMP_BASE_DIR, "expired_test.docx")
    with open(temp_file, "w") as f:
        f.write("dummy content")
        
    # Put in mapping with creation time 40 minutes ago
    file_mappings[dummy_id] = (temp_file, "expired_test.docx", time.time() - 2400)
    
    # Trigger a health check or redact call to run background prune
    # We will trigger upload to see if prune gets called, or call prune directly
    from backend.main import prune_expired_files
    prune_expired_files(max_age_seconds=1800)
    
    # Verify file and mapping are cleaned up
    assert dummy_id not in file_mappings
    assert not os.path.exists(temp_file)
