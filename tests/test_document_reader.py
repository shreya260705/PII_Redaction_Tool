import pytest
import docx
from pathlib import Path
from src.document_reader import DocumentReader


def create_test_docx(
    path: Path,
    paragraphs: list = None,
    tables: list = None,
    header_text: str = None,
    footer_text: str = None
) -> None:
    """Helper to create a temporary DOCX file with specified content."""
    doc = docx.Document()

    # Configure header and footer if provided
    if header_text or footer_text:
        section = doc.sections[0]
        # In python-docx, sections are initialized with a default header/footer containing 1 empty paragraph
        if header_text:
            section.header.paragraphs[0].text = header_text
        if footer_text:
            section.footer.paragraphs[0].text = footer_text

    # Add body paragraphs
    if paragraphs:
        for p_text in paragraphs:
            doc.add_paragraph(p_text)

    # Add body tables
    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else 0
            if rows > 0 and cols > 0:
                t = doc.add_table(rows=rows, cols=cols)
                for r_idx, row_data in enumerate(table_data):
                    for c_idx, cell_value in enumerate(row_data):
                        t.rows[r_idx].cells[c_idx].text = cell_value

    doc.save(path)


def test_read_simple_paragraphs(tmp_path):
    """Test reading a simple DOCX file containing paragraphs only."""
    file_path = tmp_path / "simple.docx"
    paras = ["First paragraph", "Second paragraph", "Third paragraph"]
    create_test_docx(file_path, paragraphs=paras)

    result = DocumentReader.read(file_path)

    assert result.metadata.source_file == "simple.docx"
    assert result.metadata.paragraph_count == 3
    assert result.metadata.table_count == 0
    assert result.metadata.table_cell_count == 0
    # There will be default empty headers/footers in the section structure,
    # but let's check that our main body paragraph blocks match exactly.
    body_blocks = [b for b in result.blocks if b.location.part_type == "body"]
    assert len(body_blocks) == 3
    for i, p_text in enumerate(paras):
        assert body_blocks[i].block_type == "paragraph"
        assert body_blocks[i].text == p_text
        assert body_blocks[i].location.paragraph_index == i
        assert body_blocks[i].location.part_type == "body"
        # Verify the element reference is preserved
        assert isinstance(body_blocks[i].element, docx.text.paragraph.Paragraph)


def test_read_table(tmp_path):
    """Test reading a DOCX containing a table."""
    file_path = tmp_path / "table.docx"
    table_data = [
        ["Header A", "Header B"],
        ["Val A1", "Val B1"],
        ["Val A2", "Val B2"]
    ]
    create_test_docx(file_path, tables=[table_data])

    result = DocumentReader.read(file_path)

    assert result.metadata.table_count == 1
    assert result.metadata.table_cell_count == 6

    table_blocks = [b for b in result.blocks if b.block_type == "table_cell"]
    assert len(table_blocks) == 6

    # Verify cells structure and ordering
    flat_cells = [cell for row in table_data for cell in row]
    for idx, cell_text in enumerate(flat_cells):
        block = table_blocks[idx]
        assert block.text == cell_text
        assert block.location.table_index == 0
        assert block.location.row_index == idx // 2
        assert block.location.cell_index == idx % 2
        assert isinstance(block.element, docx.table._Cell)


def test_read_empty_paragraph(tmp_path):
    """Test that empty paragraphs are preserved and extracted as empty text blocks."""
    file_path = tmp_path / "empty.docx"
    paras = ["Para 1", "", "Para 2"]
    create_test_docx(file_path, paragraphs=paras)

    result = DocumentReader.read(file_path)
    body_blocks = [b for b in result.blocks if b.location.part_type == "body"]
    
    assert len(body_blocks) == 3
    assert body_blocks[1].text == ""
    assert body_blocks[1].block_type == "paragraph"
    assert body_blocks[1].location.paragraph_index == 1


def test_read_headers_footers(tmp_path):
    """Test extraction of headers and footers from document sections."""
    file_path = tmp_path / "header_footer.docx"
    create_test_docx(file_path, paragraphs=["Body text"], header_text="Test Header", footer_text="Test Footer")

    result = DocumentReader.read(file_path)

    header_blocks = [b for b in result.blocks if b.location.part_type == "header"]
    footer_blocks = [b for b in result.blocks if b.location.part_type == "footer"]

    # At least the default header and footer should have text
    assert len(header_blocks) >= 1
    assert any(b.text == "Test Header" for b in header_blocks)
    assert any(b.text == "Test Footer" for b in footer_blocks)

    # Check structural indices of the matched header
    matching_header = next(b for b in header_blocks if b.text == "Test Header")
    assert matching_header.location.section_index == 0
    assert matching_header.location.header_footer_type == "default"
    assert matching_header.location.paragraph_index == 0


def test_read_missing_file():
    """Verify FileNotFoundError is raised for a missing file."""
    with pytest.raises(FileNotFoundError):
        DocumentReader.read("non_existent_file_xyz_123.docx")


def test_read_invalid_file_extension(tmp_path):
    """Verify ValueError is raised for a file with a non-docx extension."""
    invalid_file = tmp_path / "invalid.txt"
    invalid_file.write_text("Some text")

    with pytest.raises(ValueError, match="Only '.docx' files are supported"):
        DocumentReader.read(invalid_file)


def test_read_corrupt_file(tmp_path):
    """Verify ValueError is raised for a corrupted or non-DOCX package renamed to .docx."""
    corrupt_file = tmp_path / "corrupt.docx"
    corrupt_file.write_bytes(b"This is not a zip or xml document")

    with pytest.raises(ValueError, match="Unreadable or corrupted Word document"):
        DocumentReader.read(corrupt_file)


def test_metadata_correctness(tmp_path):
    """Test all metadata totals correspond to the created elements."""
    file_path = tmp_path / "stats.docx"
    paras = ["P1", "P2", "P3", "P4"]
    table_data = [["A", "B", "C"], ["D", "E", "F"]]  # 2x3 table = 6 cells
    create_test_docx(file_path, paragraphs=paras, tables=[table_data], header_text="H", footer_text="F")

    result = DocumentReader.read(file_path)

    assert result.metadata.paragraph_count == 4
    assert result.metadata.table_count == 1
    assert result.metadata.table_cell_count == 6
    assert result.metadata.header_footer_block_count >= 2
    # Verify exact sum of all block lengths
    expected_len = sum(len(b.text) for b in result.blocks)
    assert result.metadata.total_text_length == expected_len
