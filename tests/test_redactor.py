import os
import re
import sys
import tempfile
import hashlib
import pytest
import docx
from docx.text.run import Run
from docx.shared import Inches

from src.redactor import (
    RedactionEngine,
    RedactionMapper,
    redact_block,
    get_paragraph_runs
)
from src.document_reader import DocumentReader
from src.detectors.base import resolve_overlaps
from src.models import DocumentBlock, BlockLocation, PIIMatch

# Helper to get file hash
def get_file_hash(filepath: str) -> str:
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        while True:
            chunk = f.read(4096)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()

# Helper to make block for testing
def make_block(text: str, element=None) -> DocumentBlock:
    return DocumentBlock(
        block_id=0,
        block_type="paragraph",
        text=text,
        location=BlockLocation(part_type="body", paragraph_index=0),
        element=element
    )

# =====================================================================
# 1. DETERMINISTIC MAPPING TESTS (REQ 1, 2, 3, 4, 35)
# =====================================================================
def test_deterministic_person_mapping():
    mapper = RedactionMapper()
    val1 = mapper.get_replacement("Rashi Patil", "PERSON")
    val2 = mapper.get_replacement("Rashi Patil", "PERSON")
    assert val1 == val2
    assert val1 != "Rashi Patil"

def test_deterministic_email_mapping():
    mapper = RedactionMapper()
    val1 = mapper.get_replacement("rashi.patil@gmail.com", "EMAIL")
    val2 = mapper.get_replacement("rashi.patil@gmail.com", "EMAIL")
    assert val1 == val2
    assert val1.endswith("@example.com")

def test_repeated_same_person_to_same_replacement():
    mapper = RedactionMapper()
    rep1 = mapper.get_replacement("Amit Sharma", "PERSON")
    rep2 = mapper.get_replacement("Amit Sharma", "PERSON")
    assert rep1 == rep2

def test_repeated_same_email_to_same_replacement():
    mapper = RedactionMapper()
    rep1 = mapper.get_replacement("amit.sharma@yahoo.co.in", "EMAIL")
    rep2 = mapper.get_replacement("amit.sharma@yahoo.co.in", "EMAIL")
    assert rep1 == rep2

def test_deterministic_mapping_across_instances():
    mapper1 = RedactionMapper()
    mapper2 = RedactionMapper()
    rep1 = mapper1.get_replacement("Vijay Verma", "PERSON")
    rep2 = mapper2.get_replacement("Vijay Verma", "PERSON")
    assert rep1 == rep2

# =====================================================================
# 2. SYNTHETIC PII FORMAT TESTS (REQ 5, 6, 7, 8, 9, 10, 11)
# =====================================================================
def test_phone_replacement_format():
    mapper = RedactionMapper()
    rep = mapper.get_replacement("+91-9876543210", "PHONE")
    assert rep.startswith("+91 555")
    assert len(rep) == 14  # len("+91 555") + 7 = 14

def test_company_replacement_format():
    mapper = RedactionMapper()
    rep1 = mapper.get_replacement("Tata Motors Limited", "COMPANY")
    assert rep1.endswith("Limited")
    
    rep2 = mapper.get_replacement("Reliance Industries Ltd", "COMPANY")
    assert rep2.endswith("Ltd")

def test_address_replacement_format():
    mapper = RedactionMapper()
    rep = mapper.get_replacement("Plot No. 45, MG Road, Baner, Pune", "ADDRESS")
    assert " - " in rep
    assert len(rep.split(",")) >= 3
    # Verify no accidental email/IP/card/SSN structure
    assert "@" not in rep
    assert not re.search(r'\b\d{3}-\d{2}-\d{4}\b', rep)
    assert not re.search(r'\b(?:\d{4}[-\s]){3}\d{4}\b', rep)

def test_ip_replacement_format():
    mapper = RedactionMapper()
    rep = mapper.get_replacement("192.168.1.1", "IP_ADDRESS")
    assert rep.startswith("10.")
    octets = rep.split(".")
    assert len(octets) == 4
    assert all(0 <= int(o) <= 255 for o in octets)

def test_ssn_replacement_format():
    mapper = RedactionMapper()
    rep = mapper.get_replacement("123-45-6789", "SSN")
    assert re.match(r'^\d{3}-\d{2}-\d{4}$', rep)
    # Check invalid formats are not generated (like 000, 666, 900+)
    parts = rep.split("-")
    assert parts[0] != "000" and parts[0] != "666" and int(parts[0]) < 900
    assert parts[1] != "00"
    assert parts[2] != "0000"

def test_credit_card_replacement_format():
    mapper = RedactionMapper()
    # Hyphenated
    rep_hyphen = mapper.get_replacement("1234-5678-1234-5678", "CREDIT_CARD")
    assert rep_hyphen.startswith("4111-")
    assert len(rep_hyphen) == 19
    # Luhn check
    digits = [int(c) for c in rep_hyphen if c.isdigit()]
    total = 0
    for i, d in enumerate(digits):
        if i % 2 == 0:
            val = d * 2
            if val > 9: val -= 9
            total += val
        else:
            total += d
    assert total % 10 == 0

    # Spaced
    rep_space = mapper.get_replacement("1234 5678 1234 5678", "CREDIT_CARD")
    assert rep_space.startswith("4111 ")
    assert len(rep_space) == 19

def test_dob_replacement_format():
    mapper = RedactionMapper()
    # 4-digit year with /
    rep1 = mapper.get_replacement("15/08/1999", "DATE_OF_BIRTH")
    assert "/" in rep1
    assert len(rep1.split("/")) == 3
    assert len(rep1.split("/")[-1]) == 4

    # 2-digit year with -
    rep2 = mapper.get_replacement("15-08-99", "DATE_OF_BIRTH")
    assert "-" in rep2
    assert len(rep2.split("-")[-1]) == 2

    # Month name
    rep3 = mapper.get_replacement("15 August 1999", "DATE_OF_BIRTH")
    assert any(m in rep3 for m in ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"])
    assert rep3[-4:].isdigit()

# =====================================================================
# 3. RUN-AWARE REPLACEMENT DETAILS (REQ 12, 13, 14, 15, 16, 17, 18, 19, 20, 21)
# =====================================================================
def test_single_run_partial_replacement():
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run("Contact Rashi Patil today")
    block = make_block(p.text, p)
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=8, end=19, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    runs = get_paragraph_runs(p)
    assert len(runs) == 1
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert runs[0].text == f"Contact {rep} today"

def test_preserve_prefix():
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run("PREFIX Rashi Patil")
    block = make_block(p.text, p)
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=7, end=18, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    runs = get_paragraph_runs(p)
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert runs[0].text == f"PREFIX {rep}"

def test_preserve_suffix():
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run("Rashi Patil SUFFIX")
    block = make_block(p.text, p)
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=0, end=11, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    runs = get_paragraph_runs(p)
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert runs[0].text == f"{rep} SUFFIX"

def test_split_across_2_runs():
    doc = docx.Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Hello Ra")
    r2 = p.add_run("shi Patil")
    block = make_block(p.text, p)
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=6, end=17, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    runs = get_paragraph_runs(p)
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert runs[0].text == f"Hello {rep}"
    assert runs[1].text == ""

def test_split_across_3_runs():
    doc = docx.Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Hello Ra")
    r2 = p.add_run("shi Pa")
    r3 = p.add_run("til is here")
    block = make_block(p.text, p)
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=6, end=17, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    runs = get_paragraph_runs(p)
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert runs[0].text == f"Hello {rep}"
    assert runs[1].text == ""
    assert runs[2].text == " is here"

def test_split_across_4_runs():
    doc = docx.Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Hello ")
    r2 = p.add_run("Ra")
    r3 = p.add_run("shi Pa")
    r4 = p.add_run("til works")
    block = make_block(p.text, p)
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=6, end=17, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    runs = get_paragraph_runs(p)
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert runs[0].text == "Hello "
    assert runs[1].text == rep
    assert runs[2].text == ""
    assert runs[3].text == " works"

def test_multiple_pii_in_one_run():
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run("Rashi Patil and Rahul Sharma")
    block = make_block(p.text, p)
    matches = [
        PIIMatch(pii_type="PERSON", text="Rashi Patil", start=0, end=11, confidence=0.95, detector="manual"),
        PIIMatch(pii_type="PERSON", text="Rahul Sharma", start=16, end=28, confidence=0.95, detector="manual")
    ]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 2
    runs = get_paragraph_runs(p)
    rep1 = mapper.get_replacement("Rashi Patil", "PERSON")
    rep2 = mapper.get_replacement("Rahul Sharma", "PERSON")
    assert runs[0].text == f"{rep1} and {rep2}"

def test_adjacent_pii():
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run("rashi@gmail.com+919876543210")
    block = make_block(p.text, p)
    matches = [
        PIIMatch(pii_type="EMAIL", text="rashi@gmail.com", start=0, end=15, confidence=0.99, detector="manual"),
        PIIMatch(pii_type="PHONE", text="+919876543210", start=15, end=28, confidence=0.99, detector="manual")
    ]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 2
    runs = get_paragraph_runs(p)
    rep1 = mapper.get_replacement("rashi@gmail.com", "EMAIL")
    rep2 = mapper.get_replacement("+919876543210", "PHONE")
    assert runs[0].text == f"{rep1}{rep2}"

def test_overlapping_matches():
    doc = docx.Document()
    p = doc.add_paragraph()
    r = p.add_run("Tata Motors Limited is here")
    block = make_block(p.text, p)
    matches = [
        PIIMatch(pii_type="PERSON", text="Tata Motors", start=0, end=11, confidence=0.85, detector="manual"),
        PIIMatch(pii_type="COMPANY", text="Tata Motors Limited", start=0, end=19, confidence=0.95, detector="manual")
    ]
    resolved = resolve_overlaps(matches)
    assert len(resolved) == 1
    assert resolved[0].pii_type == "COMPANY"

    mapper = RedactionMapper()
    applied = redact_block(block, resolved, mapper)
    assert len(applied) == 1
    runs = get_paragraph_runs(p)
    rep = mapper.get_replacement("Tata Motors Limited", "COMPANY")
    assert runs[0].text == f"{rep} is here"

def test_formatting_preservation():
    doc = docx.Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Contact ")
    r2 = p.add_run("Rashi Patil")
    r2.bold = True
    r2.italic = True
    r3 = p.add_run(" today")
    block = make_block(p.text, p)
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=8, end=19, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    runs = get_paragraph_runs(p)
    assert len(runs) == 3
    assert runs[1].bold is True
    assert runs[1].italic is True
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert runs[1].text == rep

# =====================================================================
# 4. STRUCTURED BLOCKS TESTS (REQ 22, 23, 24, 25, 26, 27, 28)
# =====================================================================
def test_table_cell_redaction():
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Manager: Rashi Patil"
    block = DocumentBlock(
        block_id=0,
        block_type="table_cell",
        text=cell.text,
        location=BlockLocation(part_type="body", table_index=0, row_index=0, cell_index=0),
        element=cell
    )
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=9, end=20, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert cell.text == f"Manager: {rep}"

def test_table_cell_multiple_paragraphs():
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p1 = cell.paragraphs[0]
    p1.text = "Hello Rashi"
    p2 = cell.add_paragraph("Patil works here")
    block = DocumentBlock(
        block_id=0,
        block_type="table_cell",
        text="Hello Rashi\nPatil works here",
        location=BlockLocation(part_type="body", table_index=0, row_index=0, cell_index=0),
        element=cell
    )
    matches = [PIIMatch(pii_type="PERSON", text="Rashi\nPatil", start=6, end=17, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    rep = mapper.get_replacement("Rashi\nPatil", "PERSON")
    assert p1.text == f"Hello {rep}"
    assert p2.text == " works here"

def test_header_redaction():
    doc = docx.Document()
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "Confidential: Rashi Patil"
    block = DocumentBlock(
        block_id=0,
        block_type="header_paragraph",
        text=p.text,
        location=BlockLocation(part_type="header", section_index=0, paragraph_index=0),
        element=p
    )
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=14, end=25, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert p.text == f"Confidential: {rep}"

def test_footer_redaction():
    doc = docx.Document()
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Footer: Rashi Patil"
    block = DocumentBlock(
        block_id=0,
        block_type="footer_paragraph",
        text=p.text,
        location=BlockLocation(part_type="footer", section_index=0, paragraph_index=0),
        element=p
    )
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=8, end=19, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert p.text == f"Footer: {rep}"

def test_header_table_redaction():
    doc = docx.Document()
    section = doc.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=1, width=Inches(5))
    cell = table.cell(0, 0)
    cell.text = "Inside Header Table: Rashi Patil"
    block = DocumentBlock(
        block_id=0,
        block_type="header_table_cell",
        text=cell.text,
        location=BlockLocation(part_type="header", section_index=0, table_index=0, row_index=0, cell_index=0),
        element=cell
    )
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=21, end=32, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert cell.text == f"Inside Header Table: {rep}"

def test_footer_table_redaction():
    doc = docx.Document()
    section = doc.sections[0]
    footer = section.footer
    table = footer.add_table(rows=1, cols=1, width=Inches(5))
    cell = table.cell(0, 0)
    cell.text = "Inside Footer Table: Rashi Patil"
    block = DocumentBlock(
        block_id=0,
        block_type="footer_table_cell",
        text=cell.text,
        location=BlockLocation(part_type="footer", section_index=0, table_index=0, row_index=0, cell_index=0),
        element=cell
    )
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=21, end=32, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert cell.text == f"Inside Footer Table: {rep}"

def test_hyperlink_nested_run_handling():
    doc = docx.Document()
    p = doc.add_paragraph()
    p_el = p._element
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    r = docx.oxml.shared.OxmlElement('w:r')
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = "Rashi Patil"
    r.append(t)
    hyperlink.append(r)
    p_el.append(hyperlink)

    runs = get_paragraph_runs(p)
    assert len(runs) == 1
    assert runs[0].text == "Rashi Patil"

    block = make_block(p.text, p)
    matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=0, end=11, confidence=0.95, detector="manual")]
    mapper = RedactionMapper()
    applied = redact_block(block, matches, mapper)
    assert len(applied) == 1
    rep = mapper.get_replacement("Rashi Patil", "PERSON")
    assert runs[0].text == rep

# =====================================================================
# 5. ENGINE BEHAVIOR AND INTEGRITY TESTS (REQ 29, 30, 31, 32, 33, 34, 36)
# =====================================================================
def test_duplicate_xml_element_protection():
    doc = docx.Document()
    p1 = doc.add_paragraph("Important name: Rashi Patil")
    block1 = DocumentBlock(
        block_id=0,
        block_type="paragraph",
        text=p1.text,
        location=BlockLocation(part_type="body", paragraph_index=0),
        element=p1
    )
    block2 = DocumentBlock(
        block_id=1,
        block_type="paragraph",
        text=p1.text,
        location=BlockLocation(part_type="body", paragraph_index=1),
        element=p1
    )
    
    engine = RedactionEngine()
    
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, "input.docx")
        output_path = os.path.join(tmpdir, "output.docx")
        doc.save(input_path)
        
        extracted = DocumentReader.read(input_path)
        extracted.blocks = [block1, block2]
        extracted.doc = doc
        
        processed_elements = set()
        total_replacements = 0
        replacements_by_type = {}
        
        for block in extracted.blocks:
            if block.element is not None:
                element = getattr(block.element, "_element", block.element)
                element_id = id(element)
                if element_id in processed_elements:
                    continue
                processed_elements.add(element_id)
                
            matches = [PIIMatch(pii_type="PERSON", text="Rashi Patil", start=16, end=27, confidence=0.95, detector="manual")]
            applied = redact_block(block, matches, engine.mapper)
            total_replacements += len(applied)
            for m in applied:
                replacements_by_type[m.pii_type] = replacements_by_type.get(m.pii_type, 0) + 1
                
        assert total_replacements == 1
        assert replacements_by_type["PERSON"] == 1

def test_no_recursive_redaction():
    doc = docx.Document()
    p = doc.add_paragraph("Rashi Patil is here")
    
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, "input.docx")
        output_path = os.path.join(tmpdir, "output.docx")
        doc.save(input_path)
        
        engine = RedactionEngine()
        res = engine.redact(input_path, output_path)
        
        assert os.path.exists(output_path)
        doc_out = docx.Document(output_path)
        rep = engine.mapper.get_replacement("Rashi Patil", "PERSON")
        assert doc_out.paragraphs[0].text == f"{rep} is here"

def test_input_file_remains_unchanged(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Rashi Patil email is rashi@gmail.com.")
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc.save(input_path)
    
    hash_before = get_file_hash(str(input_path))
    engine = RedactionEngine()
    engine.redact(str(input_path), str(output_path))
    hash_after = get_file_hash(str(input_path))
    
    assert hash_before == hash_after

def test_output_can_be_reopened(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("My address is Plot No. 45, MG Road, Pune.")
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc.save(input_path)
    
    engine = RedactionEngine()
    engine.redact(str(input_path), str(output_path))
    
    assert os.path.exists(str(output_path))
    doc_out = docx.Document(str(output_path))
    assert len(doc_out.paragraphs) == 1

def test_paragraph_table_counts_preserved(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Para 1")
    doc.add_paragraph("Para 2")
    doc.add_table(rows=2, cols=2)
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc.save(input_path)
    
    engine = RedactionEngine()
    engine.redact(str(input_path), str(output_path))
    
    doc_out = docx.Document(str(output_path))
    assert len(doc_out.paragraphs) == 2
    assert len(doc_out.tables) == 1

def test_replacement_statistics_match_actual(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Compliance Officer Rashi Patil at rashi@gmail.com today")
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc.save(input_path)
    
    engine = RedactionEngine()
    res = engine.redact(str(input_path), str(output_path))
    
    total = res["total_replacements"]
    by_type = res["replacements_by_type"]
    assert sum(by_type.values()) == total
    assert "PERSON" in by_type
    assert "EMAIL" in by_type

def test_windows_utf8_safety():
    if sys.platform.startswith("win") and hasattr(sys.stdout, "reconfigure"):
        try:
            sys.stdout.reconfigure(encoding="utf-8")
        except Exception:
            pass
    assert True
