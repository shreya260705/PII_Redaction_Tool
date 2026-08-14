import sys
import os
sys.path.insert(0, os.path.abspath("."))

import time
import requests
import docx
from src.document_reader import DocumentReader

VERCEL_ORIGIN = "https://pii-redaction-tool.vercel.app"
RENDER_URL = "https://pii-redaction-tool-fbh6.onrender.com"
FILE_PATH = "Red Herring Prospectus.docx"

def test_full_vercel_render_user_flow():
    print(f"=== TESTING REAL VERCEL FRONTEND -> RENDER BACKEND USER FLOW ===")
    print(f"Frontend Origin: {VERCEL_ORIGIN}")
    print(f"Backend Target: {RENDER_URL}")
    print(f"Document: {FILE_PATH} ({os.path.getsize(FILE_PATH)/(1024*1024):.2f} MB)")
    
    headers = {
        "Origin": VERCEL_ORIGIN,
        "Referer": f"{VERCEL_ORIGIN}/"
    }

    # Step 0: Wake up Render backend if sleeping
    print("\n0. Waking up Render backend...")
    for attempt in range(1, 4):
        try:
            h = requests.get(f"{RENDER_URL}/api/health", timeout=120)
            if h.status_code == 200:
                print(f"Render is awake! ({h.json()})")
                break
        except Exception as e:
            print(f"Wakeup attempt {attempt} timed out/failed: {e}. Retrying in 5s...")
            time.sleep(5)

    # Step 1: Preflight OPTIONS check
    print("\n1. Testing Browser Preflight OPTIONS /api/redact-async...")
    preflight_headers = {
        "Origin": VERCEL_ORIGIN,
        "Access-Control-Request-Method": "POST",
        "Access-Control-Request-Headers": "content-type"
    }
    opt_res = requests.options(f"{RENDER_URL}/api/redact-async", headers=preflight_headers, timeout=60)
    print(f"Preflight Status: {opt_res.status_code}")
    print(f"Access-Control-Allow-Origin: {opt_res.headers.get('access-control-allow-origin')}")
    assert opt_res.status_code in [200, 204], "Preflight request failed!"
    assert opt_res.headers.get('access-control-allow-origin') in ["*", VERCEL_ORIGIN]

    # Step 2: User drops/uploads file via frontend API call
    print("\n2. Submitting Async Upload to Render backend...")
    upload_start = time.time()
    with open(FILE_PATH, "rb") as f:
        files = {"file": (os.path.basename(FILE_PATH), f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        res = requests.post(f"{RENDER_URL}/api/redact-async", files=files, headers=headers, timeout=120)

    print(f"Upload Status Code: {res.status_code}")
    assert res.status_code == 200, f"Upload failed: {res.text}"
    
    upload_data = res.json()
    task_id = upload_data.get("task_id")
    print(f"Task Created! Task ID: {task_id}")
    assert task_id, "No task_id returned!"

    # Step 3: Frontend Polling
    print("\n3. Frontend Polling /api/tasks/{task_id} every 3s...")
    poll_start = time.time()
    task_result = None
    poll_count = 0
    
    while time.time() - poll_start < 360: # 6 mins max
        time.sleep(3)
        poll_count += 1
        try:
            poll_res = requests.get(f"{RENDER_URL}/api/tasks/{task_id}", headers=headers, timeout=30)
            if poll_res.status_code == 200:
                task_data = poll_res.json()
                status = task_data.get("status")
                print(f"  [Poll #{poll_count}] Status: {status} ({time.time() - poll_start:.1f}s)")
                if status in ["success", "error"]:
                    task_result = task_data
                    break
            else:
                print(f"  [Poll #{poll_count}] Non-200 status: {poll_res.status_code}")
        except Exception as err:
            print(f"  [Poll #{poll_count}] Network hiccup during poll: {err} (Retrying on next interval)")

    total_processing_time = time.time() - upload_start
    print(f"\nProcessing Completed in {total_processing_time:.2f}s!")
    assert task_result and task_result.get("status") == "success", f"Task failed: {task_result}"

    result_info = task_result["result"]
    file_id = result_info["file_id"]
    filename = result_info["filename"]
    replacements = result_info["total_replacements"]
    types_breakdown = result_info["replacements_by_type"]
    expires_at = result_info["expires_at"]

    print(f"\n4. Redaction Results Summary:")
    print(f"  File ID: {file_id}")
    print(f"  Download Filename: {filename}")
    print(f"  Expires At: {expires_at}")
    print(f"  Total Replacements: {replacements}")
    print(f"  Replacements Breakdown: {types_breakdown}")
    assert replacements == 1586, f"Expected 1586 replacements, got {replacements}"

    # Step 4: User clicks Download button
    print("\n5. User clicking Download button -> GET /api/download/{file_id}...")
    dl_res = requests.get(f"{RENDER_URL}/api/download/{file_id}", headers=headers, timeout=60)
    print(f"Download Response Status: {dl_res.status_code}")
    print(f"Content-Type: {dl_res.headers.get('content-type')}")
    print(f"Content-Disposition: {dl_res.headers.get('content-disposition')}")
    print(f"Downloaded Size: {len(dl_res.content)} bytes")
    assert dl_res.status_code == 200, "Download failed!"

    # Save and verify file validity
    output_filename = "output_vercel_render_user_download.docx"
    with open(output_filename, "wb") as f:
        f.write(dl_res.content)

    doc = DocumentReader.read(output_filename)
    print(f"\n6. Downloaded DOCX Integrity Check:")
    print(f"  Extracted Document Blocks: {len(doc.blocks)}")
    assert len(doc.blocks) == 4384, f"Expected 4384 blocks, got {len(doc.blocks)}"
    
    doc_parsed = docx.Document(output_filename)
    print(f"  Parsed Paragraphs: {len(doc_parsed.paragraphs)}")
    print(f"  Parsed Tables: {len(doc_parsed.tables)}")
    print("\nVERCEL -> RENDER END-TO-END USER FLOW PASSED PERFECTLY!")

if __name__ == "__main__":
    test_full_vercel_render_user_flow()
