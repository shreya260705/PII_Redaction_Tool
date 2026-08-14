import time
import docx
from docx.table import _Cell

large_doc = r"c:\Users\shrey\OneDrive\Desktop\PII-Redaction-Tool\Red Herring Prospectus.docx"
doc = docx.Document(large_doc)

print("Starting comparison of Table Cell reading methods...")

# Method 1: Original using row.cells
start_1 = time.time()
cell_count_1 = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            t = cell.text
            cell_count_1 += 1
duration_1 = time.time() - start_1
print(f"Method 1 (row.cells): Read {cell_count_1} cells in {duration_1:.2f} seconds.")

# Method 2: Optimized using row._tr.tc_lst
start_2 = time.time()
cell_count_2 = 0
for table in doc.tables:
    for row in table.rows:
        for tc in row._tr.tc_lst:
            cell = _Cell(tc, table)
            t = cell.text
            cell_count_2 += 1
duration_2 = time.time() - start_2
print(f"Method 2 (row._tr.tc_lst): Read {cell_count_2} cells in {duration_2:.2f} seconds.")
