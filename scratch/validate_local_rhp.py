import sys
import os
sys.path.insert(0, os.path.abspath("."))

import time
import docx
from src.redactor import RedactionEngine
from src.document_reader import DocumentReader

def main():
    input_path = "Red Herring Prospectus.docx"
    output_path = "output_redacted_large.docx"

    start_time = time.time()
    engine = RedactionEngine()
    results = engine.redact(input_path, output_path)
    elapsed = time.time() - start_time

    out_size = os.path.getsize(output_path) / (1024 * 1024)

    print("=== RHP REDACTION LOCAL PIPELINE RESULTS ===")
    print(f"Processing Time: {elapsed:.2f} seconds")
    print(f"Total Replacements: {results['total_replacements']}")
    print(f"Replacements by Type:")
    for k, v in sorted(results['replacements_by_type'].items()):
        print(f"  {k:<15}: {v}")
    print(f"Output File Size: {out_size:.2f} MB")

    # Verify DOCX validity & readability
    doc = DocumentReader.read(output_path)
    print(f"Redacted Document Readability: OK ({len(doc.blocks)} blocks extracted)")

    # Verify output docx can be opened natively by python-docx
    doc_docx = docx.Document(output_path)
    print(f"docx.Document parse check: OK ({len(doc_docx.paragraphs)} paragraphs, {len(doc_docx.tables)} tables)")

if __name__ == "__main__":
    main()
